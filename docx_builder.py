"""Word export logic for the Presentation PowerPoint Builder.

The mentor review document is the review workspace. The app does not store mentor
critiques; mentors can use Word comments or Track Changes in the generated DOCX.
"""

from __future__ import annotations

import base64
from io import BytesIO
from typing import Any, Dict, List
from zipfile import ZIP_DEFLATED, ZipFile

from PIL import Image

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from deck_model import APP_VERSION, BLOOM_HELPER, identity_subtitle, identity_title, slide_output_title, split_nonempty_lines
from preview_utils import render_pptx_first_slide_to_png

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
HEADER_GRAY = "D9D9D9"
PALE_GRAY = "F7F7F7"
PINK = "F4CCCC"
WHITE = "FFFFFF"
BORDER = "000000"
DOC_FONT = "Calibri"
TEXT_DARK = RGBColor(25, 25, 25)
TEXT_MUTED = RGBColor(95, 95, 95)
EMU_PER_INCH = 914400
TWIPS_PER_INCH = 1440


# -----------------------------------------------------------------------------
# Low-level table formatting helpers
# -----------------------------------------------------------------------------


def _safe_text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _visual_image_bytes(slide: Dict[str, Any]) -> bytes | None:
    image = slide.get("visual_image", {})
    if not isinstance(image, dict):
        return None
    encoded = image.get("data_base64")
    if not encoded:
        return None
    try:
        return base64.b64decode(encoded)
    except Exception:
        return None


def _visual_image_filename(slide: Dict[str, Any]) -> str:
    image = slide.get("visual_image", {})
    if not isinstance(image, dict):
        return ""
    return _safe_text(image.get("filename"))


def _uploaded_slide_pptx_filename(slide: Dict[str, Any]) -> str:
    pptx = slide.get("uploaded_slide_pptx", {})
    if not isinstance(pptx, dict):
        return ""
    return _safe_text(pptx.get("filename"))


def _uploaded_slide_pptx_bytes(slide: Dict[str, Any]) -> bytes | None:
    pptx = slide.get("uploaded_slide_pptx", {})
    if not isinstance(pptx, dict):
        return None
    encoded = pptx.get("data_base64")
    if not encoded:
        return None
    try:
        return base64.b64decode(encoded)
    except Exception:
        return None


def _uploaded_slide_preview_image_bytes(slide: Dict[str, Any]) -> bytes | None:
    image = slide.get("uploaded_slide_preview_image", {})
    if not isinstance(image, dict):
        return None
    encoded = image.get("data_base64")
    if not encoded:
        return None
    try:
        return base64.b64decode(encoded)
    except Exception:
        return None


def _ensure_uploaded_slide_preview_image(slide: Dict[str, Any]) -> bytes | None:
    preview = _uploaded_slide_preview_image_bytes(slide)
    if preview:
        return preview
    pptx_bytes = _uploaded_slide_pptx_bytes(slide)
    if not pptx_bytes:
        return None
    preview = render_pptx_first_slide_to_png(pptx_bytes)
    if preview:
        filename = _uploaded_slide_pptx_filename(slide) or "uploaded_slide.pptx"
        stem = filename.rsplit('.', 1)[0]
        slide["uploaded_slide_preview_image"] = {
            "filename": f"{stem}_preview.png",
            "content_type": "image/png",
            "data_base64": base64.b64encode(preview).decode("ascii"),
        }
    return preview


def _body_width_inches(doc: Document) -> float:
    section = doc.sections[-1]
    return float(section.page_width - section.left_margin - section.right_margin) / EMU_PER_INCH


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * TWIPS_PER_INCH)))
    tc_w.set(qn("w:type"), "dxa")


def _lock_table_widths(table, widths: List[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * TWIPS_PER_INCH)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                _set_cell_width(row.cells[idx], width)


def _set_row_height(row, height_pt: float, exact: bool = False) -> None:
    row.height = Pt(height_pt)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(int(height_pt * 20)))
    tr_height.set(qn("w:hRule"), "exact" if exact else "atLeast")


def _prevent_row_split(row) -> None:
    """Keep a table row together on one page when Word paginates the DOCX."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _clear_cell(cell) -> None:
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()


def _write_cell_text(
    cell,
    text: Any,
    *,
    font_size: float = 9.0,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor = TEXT_DARK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing: float = 1.0,
) -> None:
    _clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraphs = _safe_text(text).splitlines() or [""]
    for idx, raw in enumerate(paragraphs):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = line_spacing
        run = p.add_run(raw)
        run.font.name = DOC_FONT
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color


# -----------------------------------------------------------------------------
# Document building blocks
# -----------------------------------------------------------------------------


def _add_title_block(doc: Document, deck: Dict[str, Any]) -> None:
    width = _body_width_inches(doc)
    table = doc.add_table(rows=3, cols=1)
    table.style = "Table Grid"
    _lock_table_widths(table, [width])

    cell = table.cell(0, 0)
    _shade_cell(cell, BLUE)
    _set_cell_borders(cell)
    _set_cell_margins(cell, 90, 120, 90, 120)
    _write_cell_text(cell, "MENTOR POWERPOINT REVIEW DOCUMENT", font_size=14, bold=True, color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

    cell = table.cell(1, 0)
    _shade_cell(cell, LIGHT_BLUE)
    _set_cell_borders(cell)
    _set_cell_margins(cell, 90, 120, 90, 120)
    _write_cell_text(cell, f"{identity_title(deck)}\n{identity_subtitle(deck)}", font_size=10.5, bold=False, color=TEXT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)

    cell = table.cell(2, 0)
    _shade_cell(cell, PALE_GRAY)
    _set_cell_borders(cell)
    _set_cell_margins(cell, 55, 90, 55, 90)
    _write_cell_text(
        cell,
        f"Generated by Pediatric Residency Presentation Builder {APP_VERSION} - complete presentation plan, story arc, structured objectives, take-home points, visuals, and speaker notes included.",
        font_size=8.3,
        bold=True,
        color=TEXT_MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_presentation_overview(doc: Document, deck: Dict[str, Any]) -> None:
    """Add every presentation-level field from the app to the mentor document."""
    meta = deck.get("metadata", {}) if isinstance(deck, dict) else {}
    width = _body_width_inches(doc)

    _add_slide_banner(doc, "Presentation plan and story arc")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    _lock_table_widths(table, [1.65, width - 1.65])

    overview_rows = [
        ("Presentation title", _safe_text(meta.get("presentation_title"))),
        ("Presentation subtitle", _safe_text(meta.get("presentation_subtitle"))),
        ("Presenter", _safe_text(meta.get("presenter"))),
        ("Session date", _safe_text(meta.get("session_date"))),
        ("Audience", _safe_text(meta.get("audience"))),
        ("Presentation type", _safe_text(meta.get("presentation_type"))),
        ("Core question / tension", _safe_text(meta.get("core_question"))),
        ("Story arc", _safe_text(meta.get("story_arc"))),
        ("Internal archive notes", _safe_text(meta.get("archive_notes"))),
        ("App version", _safe_text(deck.get("app_version"))),
        ("Total slides", str(len(deck.get("slides", [])))),
    ]
    for label, value in overview_rows:
        _add_field_row(table, label, value)

    # Include any future nonblank metadata fields automatically.
    known_metadata_keys = {
        "presentation_title", "presentation_subtitle", "presenter",
        "session_date", "audience", "presentation_type", "core_question",
        "story_arc", "archive_notes",
    }
    for key, raw_value in meta.items():
        if key in known_metadata_keys:
            continue
        value = _safe_text(raw_value)
        if value:
            label = key.replace("_", " ").strip().title()
            _add_field_row(table, f"Additional metadata - {label}", value)

    _add_field_row(
        table,
        "Overall mentor review",
        "[Comment on the overall story, alignment between title/objectives/content, sequencing, omissions, and major revisions.]",
        PINK,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_guidelines(doc: Document) -> None:
    width = _body_width_inches(doc)
    table = doc.add_table(rows=5, cols=1)
    table.style = "Table Grid"
    _lock_table_widths(table, [width])

    rows = [
        ("How to use this document", LIGHT_BLUE, True),
        ("Use Word comments or Track Changes to critique slide wording, clarity, educational flow, accuracy, and speaker notes. Do not rewrite the presentation inside the app. The resident should return to the app and make final edits there.", WHITE, False),
        ("Mentor focus", LIGHT_BLUE, True),
        ("1. Does the title/objectives match the actual story?\n2. Do the slide titles tell a clear beginning-middle-end narrative?\n3. Is each slide readable and teachable?\n4. Are the speaker notes useful enough for rehearsal?\n5. Are results/data interpreted rather than dumped?", WHITE, False),
        ("Track Changes is enabled in this document so edits are easier to review.", PALE_GRAY, False),
    ]
    for i, (text, fill, bold) in enumerate(rows):
        cell = table.cell(i, 0)
        _shade_cell(cell, fill)
        _set_cell_borders(cell)
        _set_cell_margins(cell)
        _write_cell_text(cell, text, font_size=9.2 if not bold else 9.8, bold=bold, color=TEXT_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_bloom_reference(doc: Document) -> None:
    width = _body_width_inches(doc)
    table = doc.add_table(rows=1 + len(BLOOM_HELPER), cols=2)
    table.style = "Table Grid"
    _lock_table_widths(table, [1.35, width - 1.35])

    h1, h2 = table.rows[0].cells
    for cell, text in [(h1, "Bloom level"), (h2, "Useful verbs")]:
        _shade_cell(cell, HEADER_GRAY)
        _set_cell_borders(cell)
        _set_cell_margins(cell)
        _write_cell_text(cell, text, font_size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, (level, verbs) in enumerate(BLOOM_HELPER.items(), start=1):
        c1, c2 = table.rows[row_idx].cells
        _shade_cell(c1, WHITE)
        _shade_cell(c2, WHITE)
        _set_cell_borders(c1)
        _set_cell_borders(c2)
        _set_cell_margins(c1)
        _set_cell_margins(c2)
        _write_cell_text(c1, level, font_size=8.5, bold=True, color=RGBColor(31, 78, 121))
        _write_cell_text(c2, verbs, font_size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_slide_banner(doc: Document, text: str) -> None:
    width = _body_width_inches(doc)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _lock_table_widths(table, [width])
    cell = table.cell(0, 0)
    _shade_cell(cell, BLUE)
    _set_cell_borders(cell)
    _set_cell_margins(cell, 65, 90, 65, 90)
    _write_cell_text(cell, text, font_size=10.5, bold=True, color=RGBColor(255, 255, 255))
    _set_row_height(table.rows[0], 20, exact=False)


def _add_field_row(table, label: str, value: str, label_fill: str = HEADER_GRAY) -> None:
    row = table.add_row()
    _prevent_row_split(row)
    label_cell, value_cell = row.cells
    _shade_cell(label_cell, label_fill)
    _shade_cell(value_cell, WHITE)
    _set_cell_borders(label_cell)
    _set_cell_borders(value_cell)
    _set_cell_margins(label_cell)
    _set_cell_margins(value_cell)
    _write_cell_text(label_cell, label, font_size=8.5, bold=True, color=RGBColor(31, 78, 121))
    _write_cell_text(value_cell, value or "[blank]", font_size=9.0, italic=not bool(value), color=TEXT_MUTED if not value else TEXT_DARK)


def _fit_image_dimensions(image_bytes: bytes, max_width: float = 3.55, max_height: float = 1.60) -> tuple[float, float]:
    """Return compact image dimensions in inches for the mentor review preview cell."""
    try:
        with Image.open(BytesIO(image_bytes)) as img:
            px_w, px_h = img.size
    except Exception:
        return max_width, max_height

    if not px_w or not px_h:
        return max_width, max_height

    image_ratio = px_w / px_h
    box_ratio = max_width / max_height
    if image_ratio >= box_ratio:
        width = max_width
        height = max_width / image_ratio
    else:
        height = max_height
        width = max_height * image_ratio
    return max(0.1, width), max(0.1, height)


def _add_image_row(table, slide: Dict[str, Any], label_fill: str = HEADER_GRAY) -> None:
    pptx_name = _uploaded_slide_pptx_filename(slide)
    image_bytes = _visual_image_bytes(slide)
    image_note = ""
    if pptx_name:
        image_bytes = _ensure_uploaded_slide_preview_image(slide)
        image_note = f"Editable PPTX slide replacement uploaded: {pptx_name}. Preview shown below. The first slide is imported into the exported PowerPoint as editable objects."

    if not image_bytes and not image_note:
        _add_field_row(table, "Uploaded visual", "[none]", label_fill)
        return

    row = table.add_row()
    _prevent_row_split(row)
    label_cell, value_cell = row.cells
    _shade_cell(label_cell, label_fill)
    _shade_cell(value_cell, WHITE)
    _set_cell_borders(label_cell)
    _set_cell_borders(value_cell)
    _set_cell_margins(label_cell)
    _set_cell_margins(value_cell)
    _write_cell_text(label_cell, "Uploaded visual", font_size=8.5, bold=True, color=RGBColor(31, 78, 121))

    _clear_cell(value_cell)
    if image_note:
        note_p = value_cell.paragraphs[0]
        note_p.paragraph_format.space_after = Pt(4)
        note_run = note_p.add_run(image_note)
        note_run.font.name = DOC_FONT
        note_run.font.size = Pt(8.5)
        note_run.font.color.rgb = TEXT_MUTED
    try:
        if image_bytes:
            image_width, image_height = _fit_image_dimensions(image_bytes)
            image_p = value_cell.add_paragraph() if image_note else value_cell.paragraphs[0]
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_p.paragraph_format.space_after = Pt(0)
            image_run = image_p.add_run()
            image_run.add_picture(
                BytesIO(image_bytes),
                width=Inches(image_width),
                height=Inches(image_height),
            )
        elif image_note:
            value_cell.add_paragraph("[Preview image could not be generated for this PPTX in the current environment.]")
    except Exception:
        value_cell.add_paragraph("[Uploaded visual could not be rendered in Word preview]")


def _add_slide_review_block(doc: Document, deck: Dict[str, Any], slide: Dict[str, Any], index: int) -> None:
    """Add a complete slide-by-slide record of every editable app field."""
    title = slide_output_title(deck, slide, index)
    role = _safe_text(slide.get("role")) or "Slide"
    kind = _safe_text(slide.get("slide_kind")) or "content"
    meta = deck.get("metadata", {}) if isinstance(deck, dict) else {}
    _add_slide_banner(doc, f"Slide {index}: {title}")

    width = _body_width_inches(doc)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    _lock_table_widths(table, [1.65, width - 1.65])

    # Fields shared by all slide types.
    _add_field_row(table, "Role", role)
    _add_field_row(table, "Slide template / type", kind)
    _add_field_row(table, "Required slide", "Yes" if bool(slide.get("required", False)) else "No")
    _add_field_row(table, "Helper prompt", _safe_text(slide.get("prompt")))
    _add_field_row(table, "PowerPoint output title", title)
    _add_field_row(table, "Slide title", _safe_text(slide.get("title")))
    _add_field_row(table, "Subtitle", _safe_text(slide.get("subtitle")))
    _add_field_row(table, "Section box label", _safe_text(slide.get("section_box_label")))

    # The title slide is driven by presentation metadata rather than ordinary body fields.
    if role == "Title" or kind == "title":
        _add_field_row(table, "Presentation title", _safe_text(meta.get("presentation_title")))
        _add_field_row(table, "Presentation subtitle", _safe_text(meta.get("presentation_subtitle")))
        _add_field_row(table, "Presenter", _safe_text(meta.get("presenter")))
        _add_field_row(table, "Session date", _safe_text(meta.get("session_date")))
        _add_field_row(table, "Audience", _safe_text(meta.get("audience")))
        _add_field_row(table, "Presentation type", _safe_text(meta.get("presentation_type")))
        _add_field_row(table, "Core question / tension", _safe_text(meta.get("core_question")))
        _add_field_row(table, "Story arc", _safe_text(meta.get("story_arc")))
        _add_field_row(table, "Internal archive notes", _safe_text(meta.get("archive_notes")))

    # Objectives use structured fields in the app; show each one explicitly.
    elif role == "Objectives" or kind == "objectives":
        _add_field_row(table, "Objectives intro", _safe_text(slide.get("objectives_intro")))
        for objective_number in range(1, 4):
            _add_field_row(
                table,
                f"Objective {objective_number} action word",
                _safe_text(slide.get(f"objective_{objective_number}_verb")),
            )
            _add_field_row(
                table,
                f"Objective {objective_number} sentence",
                _safe_text(slide.get(f"objective_{objective_number}_text")),
            )
        _add_field_row(table, "Bottom takeaway banner", _safe_text(slide.get("objectives_takeaway")))
        _add_field_row(table, "Compiled slide text", _safe_text(slide.get("body")))

    # Take-home slides also use five structured fields.
    elif role == "Take-home" or kind == "takehome":
        for point_number in range(1, 6):
            _add_field_row(
                table,
                f"Take-home point {point_number}",
                _safe_text(slide.get(f"takehome_point_{point_number}")),
            )
        _add_field_row(table, "Compiled slide text", _safe_text(slide.get("body")))

    # Ordinary slides use the main body field.
    else:
        _add_field_row(table, "Slide text", _safe_text(slide.get("body")))

    # Include any future nonblank slide fields automatically, without exposing
    # binary image/PPTX payloads or duplicating fields already shown above.
    represented_keys = {
        "id", "role", "slide_kind", "required", "prompt", "title", "subtitle",
        "section_box_label", "body", "visual_plan", "visual_image",
        "uploaded_slide_pptx", "uploaded_slide_preview_image", "visual_full_slide", "discussion_prompt",
        "speaker_notes", "objectives_intro", "objective_1_verb",
        "objective_1_text", "objective_2_verb", "objective_2_text",
        "objective_3_verb", "objective_3_text", "objectives_takeaway",
        "takehome_point_1", "takehome_point_2", "takehome_point_3",
        "takehome_point_4", "takehome_point_5",
    }
    for key, raw_value in slide.items():
        if key in represented_keys:
            continue
        value = _safe_text(raw_value)
        if value:
            label = key.replace("_", " ").strip().title()
            _add_field_row(table, f"Additional slide field - {label}", value)

    # Keep any nonblank legacy visual-plan content so old saved presentations lose nothing.
    if _safe_text(slide.get("visual_plan")):
        _add_field_row(table, "Legacy visual / evidence plan", _safe_text(slide.get("visual_plan")))

    _add_image_row(table, slide)
    if _uploaded_slide_pptx_bytes(slide):
        visual_layout = "Editable PPTX slide replacement"
    elif _visual_image_bytes(slide):
        visual_layout = "Whole-slide visual" if bool(slide.get("visual_full_slide", False)) else "Image used within slide layout"
    else:
        visual_layout = "No uploaded visual"
    _add_field_row(table, "Visual layout", visual_layout)
    _add_field_row(table, "Discussion prompt", _safe_text(slide.get("discussion_prompt")))
    _add_field_row(table, "Speaker notes", _safe_text(slide.get("speaker_notes")))
    _add_field_row(table, "Mentor notes", "[Add comments here or use Word comments in the margin.]", PINK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _enable_track_changes(docx_stream: BytesIO) -> BytesIO:
    """Open the generated mentor document with Track Changes enabled."""
    try:
        source = BytesIO(docx_stream.getvalue())
        target = BytesIO()
        with ZipFile(source, "r") as zin, ZipFile(target, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "w:trackRevisions" not in xml:
                        xml = xml.replace("</w:settings>", "<w:trackRevisions/></w:settings>")
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
        target.seek(0)
        return target
    except Exception:
        docx_stream.seek(0)
        return docx_stream


def build_mentor_review_docx(deck: Dict[str, Any]) -> bytes:
    """Build an editable mentor review Word document."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = DOC_FONT
    styles["Normal"].font.size = Pt(9.5)

    _add_title_block(doc, deck)
    _add_presentation_overview(doc, deck)
    _add_guidelines(doc)
    _add_bloom_reference(doc)

    for idx, slide in enumerate(deck.get("slides", []), start=1):
        if idx > 1:
            doc.add_page_break()
        _add_slide_review_block(doc, deck, slide, idx)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    reviewed = _enable_track_changes(output)
    return reviewed.getvalue()


def mentor_docx_contains_complete_review_fields(docx_bytes: bytes) -> bool:
    """Confirm that the current complete mentor-review template was generated."""
    required_labels = [
        "Presentation plan and story arc",
        "Core question / tension",
        "Story arc",
        "Overall mentor review",
        "Speaker notes",
        f"Generated by Pediatric Residency Presentation Builder {APP_VERSION}",
    ]
    try:
        with ZipFile(BytesIO(docx_bytes), "r") as zin:
            xml = zin.read("word/document.xml").decode("utf-8", errors="ignore")
        return all(label in xml for label in required_labels)
    except Exception:
        return False


def build_plain_text_summary(deck: Dict[str, Any]) -> str:
    """Complete text representation used by tests and future integrations."""
    meta = deck.get("metadata", {}) if isinstance(deck, dict) else {}
    parts = [identity_title(deck), identity_subtitle(deck), ""]
    for key in [
        "presentation_subtitle",
        "core_question",
        "story_arc",
        "archive_notes",
    ]:
        value = _safe_text(meta.get(key))
        if value:
            parts.append(f"{key}: {value}")
    parts.append("")

    excluded = {"id", "visual_image", "uploaded_slide_pptx", "uploaded_slide_preview_image"}
    for idx, slide in enumerate(deck.get("slides", []), start=1):
        parts.append(f"Slide {idx}: {slide_output_title(deck, slide, idx)}")
        for key, raw_value in slide.items():
            if key in excluded:
                continue
            value = _safe_text(raw_value)
            if value:
                parts.append(f"{key}: {value}")
        if _visual_image_bytes(slide):
            parts.append("uploaded_visual: image")
        if _uploaded_slide_pptx_bytes(slide):
            parts.append("uploaded_visual: editable PPTX replacement")
        parts.append("")
    return "\n".join(parts)
